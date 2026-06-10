import ast
import logging
import re
from pathlib import Path

from langchain_core.documents import Document

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS: set[str] = {
    ".py", ".ts", ".tsx", ".js", ".jsx",
    ".go", ".java", ".kt", ".rb", ".rs",
    ".cpp", ".c", ".h", ".cs", ".php",
    ".md", ".txt", ".yaml", ".yml", ".toml", ".json",
    ".pdf", ".docx",
}

IGNORE_DIRS: set[str] = {
    ".git", "__pycache__", "node_modules", ".next", "dist", "build",
    "venv", ".venv", ".env", "coverage", ".pytest_cache", ".mypy_cache",
}

MAX_FILE_BYTES = 200_000  # skip files > 200 KB


def chunk_python_file(source: str, file_path: str, codebase_id: str) -> list[Document]:
    """AST-level chunking for Python: yields one Document per function/class."""
    try:
        tree = ast.parse(source)
    except SyntaxError:
        logger.warning("AST parse failed for %s — falling back to character split", file_path)
        return chunk_by_characters(source, file_path, "python", codebase_id)

    lines = source.splitlines()
    docs: list[Document] = []

    for node in ast.walk(tree):
        if not isinstance(node, (ast.FunctionDef, ast.AsyncFunctionDef, ast.ClassDef)):
            continue
        start = node.lineno - 1
        end = node.end_lineno if node.end_lineno else start + 1
        chunk = "\n".join(lines[start:end])
        if not chunk.strip():
            continue
        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "file_path": file_path,
                    "language": "python",
                    "symbol_name": node.name,
                    "start_line": start + 1,
                    "end_line": end,
                    "chunk_type": "class" if isinstance(node, ast.ClassDef) else "function",
                    "codebase_id": codebase_id,
                },
            )
        )
    if not docs:
        # file has no top-level functions/classes — store whole file
        docs = chunk_by_characters(source, file_path, "python", codebase_id)
    return docs


_TS_FUNC_RE = re.compile(
    r"(?:^|\n)((?:export\s+)?(?:async\s+)?function\s+\w+[^{]*\{)",
    re.MULTILINE,
)
_TS_CLASS_RE = re.compile(
    r"(?:^|\n)((?:export\s+)?(?:abstract\s+)?class\s+\w+[^{]*\{)",
    re.MULTILINE,
)


def chunk_ts_file(source: str, file_path: str, lang: str, codebase_id: str) -> list[Document]:
    """Heuristic chunking for TypeScript/JavaScript by function/class blocks."""
    matches = list(_TS_FUNC_RE.finditer(source)) + list(_TS_CLASS_RE.finditer(source))
    if not matches:
        return chunk_by_characters(source, file_path, lang, codebase_id)

    matches.sort(key=lambda m: m.start())
    docs: list[Document] = []
    lines = source.splitlines()

    for i, match in enumerate(matches):
        start_char = match.start()
        end_char = matches[i + 1].start() if i + 1 < len(matches) else len(source)
        chunk = source[start_char:end_char].strip()
        if not chunk:
            continue
        start_line = source[:start_char].count("\n") + 1
        end_line = start_line + chunk.count("\n")
        symbol = re.search(r"(?:function|class)\s+(\w+)", match.group(0))
        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "file_path": file_path,
                    "language": lang,
                    "symbol_name": symbol.group(1) if symbol else "",
                    "start_line": start_line,
                    "end_line": end_line,
                    "chunk_type": "class" if "class" in match.group(0) else "function",
                    "codebase_id": codebase_id,
                },
            )
        )
    return docs if docs else chunk_by_characters(source, file_path, lang, codebase_id)


def chunk_by_characters(
    source: str,
    file_path: str,
    language: str,
    codebase_id: str,
    chunk_size: int = 1000,
    overlap: int = 200,
) -> list[Document]:
    """Fallback: split by character count with overlap."""
    docs: list[Document] = []
    start = 0
    line_offset = 0
    while start < len(source):
        end = min(start + chunk_size, len(source))
        chunk = source[start:end]
        start_line = source[:start].count("\n") + 1
        end_line = start_line + chunk.count("\n")
        docs.append(
            Document(
                page_content=chunk,
                metadata={
                    "file_path": file_path,
                    "language": language,
                    "symbol_name": "",
                    "start_line": start_line,
                    "end_line": end_line,
                    "chunk_type": "generic",
                    "codebase_id": codebase_id,
                },
            )
        )
        start += chunk_size - overlap
    return docs


def _extract_pdf(path: Path) -> str:
    """Extract plain text from a PDF using pypdf."""
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]
        reader = PdfReader(str(path))
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text() or ""
            parts.append(text)
        return "\n".join(parts)
    except Exception as exc:  # noqa: BLE001
        logger.warning("PDF extraction failed for %s: %s", path, exc)
        return ""


def _extract_docx(path: Path) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        from docx import Document as DocxDocument  # type: ignore[import-untyped]
        doc = DocxDocument(str(path))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as exc:  # noqa: BLE001
        logger.warning("DOCX extraction failed for %s: %s", path, exc)
        return ""


def chunk_file(path: Path, root: Path, codebase_id: str) -> list[Document]:
    """Chunk a single file into Document objects based on its language."""
    ext = path.suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return []
    if path.stat().st_size > MAX_FILE_BYTES:
        logger.debug("Skipping large file: %s", path)
        return []

    relative_path = str(path.relative_to(root))

    # Binary document types — extract text first
    if ext == ".pdf":
        source = _extract_pdf(path)
        if not source.strip():
            logger.warning("No text extracted from PDF: %s", path)
            return []
        return chunk_by_characters(source, relative_path, "pdf", codebase_id)

    if ext == ".docx":
        source = _extract_docx(path)
        if not source.strip():
            logger.warning("No text extracted from DOCX: %s", path)
            return []
        return chunk_by_characters(source, relative_path, "docx", codebase_id)

    try:
        source = path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        logger.warning("Cannot read %s: %s", path, exc)
        return []

    if ext == ".py":
        return chunk_python_file(source, relative_path, codebase_id)
    if ext in {".ts", ".tsx"}:
        return chunk_ts_file(source, relative_path, "typescript", codebase_id)
    if ext in {".js", ".jsx"}:
        return chunk_ts_file(source, relative_path, "javascript", codebase_id)
    return chunk_by_characters(source, relative_path, ext.lstrip("."), codebase_id)


def walk_directory(root: Path, codebase_id: str) -> list[Document]:
    """Recursively walk a directory and chunk all supported files."""
    all_docs: list[Document] = []
    for path in root.rglob("*"):
        if path.is_dir():
            continue
        if any(part in IGNORE_DIRS for part in path.parts):
            continue
        docs = chunk_file(path, root, codebase_id)
        all_docs.extend(docs)
    logger.info("Chunked %d documents from %s", len(all_docs), root)
    return all_docs
